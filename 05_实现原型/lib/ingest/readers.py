"""
readers.py — 多格式文件全文读取器（v0.21.5）

职责：把单文件（md/txt/pdf/docx/csv/json）完整读为字符串，
      或批量读一个 learner 的全部文件全文。

与 library_store.py 的区别：
- library_store.read_user_file_text 是"前 N 字 + 限前 3 页"的摘要视图，用于 system prompt 注入；
- readers.read_file_full 是"完整文本"，用于后续 chunker 切片 / 检索 / 重组结构。

设计原则：
1. 缺库不崩：pypdf / python-docx 缺失时打印警告并尝试降级到文本读取
2. 失败不崩：单文件解析失败返回 {"ok": False, "text": ""}，绝不抛出
3. 与 library_store 共用一份"扩展名→提取器"语义，但去掉 800 字 / 3 页 / 50 行的限制
4. read_corpus_full 通过 library_store.list_user_files 取文件清单（双读 usr_knowledge + 旧路径）
"""
from __future__ import annotations

import csv
import json
import logging
import sys
from pathlib import Path
from typing import List

# 把上级目录（lib/）加到 sys.path，方便 library_store 直引
_LIB_PARENT = Path(__file__).resolve().parent.parent
if str(_LIB_PARENT) not in sys.path:
    sys.path.insert(0, str(_LIB_PARENT))

import library_store  # noqa: E402

log = logging.getLogger("ingest.readers")


# ────────────────────────────────────────────────────────────────
# 各格式全文提取器（无字数 / 页数 / 行数上限）
# ────────────────────────────────────────────────────────────────

def _read_text_file_full(path: Path) -> str:
    """UTF-8 → GBK → GB18030 → latin-1 兜底；按行返回完整文本。"""
    try:
        raw = path.read_bytes()
    except Exception as e:
        log.warning("[readers] 读取文件字节失败 %s: %s", path, e)
        return ""
    for enc in ("utf-8-sig", "utf-8", "gbk", "gb18030", "latin-1"):
        try:
            return raw.decode(enc, errors="ignore")
        except Exception:
            continue
    log.warning("[readers] 所有编码都失败 %s", path)
    return ""


def _read_pdf_full(path: Path) -> str:
    """pypdf 全页提取；缺库时打 warning 返回空串。"""
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        log.warning("[readers] pypdf 未安装，跳过 PDF 全文: %s", path)
        return ""
    try:
        reader = PdfReader(str(path))
        buf: list[str] = []
        for i, page in enumerate(reader.pages):
            try:
                txt = page.extract_text() or ""
            except Exception as e:
                log.warning("[readers] PDF 第 %d 页提取失败: %s", i + 1, e)
                txt = ""
            if txt:
                buf.append(txt)
        return "\n".join(buf)
    except Exception as e:
        log.warning("[readers] PDF 全页提取失败 %s: %s", path, e)
        return ""


def _read_docx_full(path: Path) -> str:
    """python-docx 全段落（含表格内文本）；缺库打 warning。"""
    try:
        import docx  # type: ignore
    except Exception:
        log.warning("[readers] python-docx 未安装，跳过 DOCX 全文: %s", path)
        return ""
    try:
        d = docx.Document(str(path))
        out: list[str] = []
        # 段落
        for para in d.paragraphs:
            t = (para.text or "").strip()
            if t:
                out.append(t)
        # 表格内文本（每行一格用 tab 分隔）
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    out.append("\t".join(cells))
        return "\n".join(out)
    except Exception as e:
        log.warning("[readers] DOCX 全文提取失败 %s: %s", path, e)
        return ""


def _read_csv_full(path: Path) -> str:
    """CSV 全部行转 tab 分隔文本（保留列结构）；utf-8-sig 优先。"""
    try:
        with open(path, "r", encoding="utf-8-sig", errors="ignore", newline="") as f:
            reader = csv.reader(f)
            buf: list[str] = []
            for row in reader:
                buf.append("\t".join(row))
            return "\n".join(buf)
    except Exception as e:
        log.warning("[readers] CSV 全文提取失败 %s: %s", path, e)
        return ""


def _read_json_full(path: Path) -> str:
    """JSON 格式化（ensure_ascii=False 保留中文）；失败降级文本读。"""
    try:
        raw = path.read_text(encoding="utf-8", errors="ignore")
        data = json.loads(raw)
        return json.dumps(data, ensure_ascii=False, indent=2)
    except json.JSONDecodeError:
        # 不是合法 JSON，按文本读
        return _read_text_file_full(path)
    except Exception as e:
        log.warning("[readers] JSON 全文提取失败 %s: %s", path, e)
        return ""


_FULL_EXTRACTORS = {
    ".md": _read_text_file_full,
    ".txt": _read_text_file_full,
    ".pdf": _read_pdf_full,
    ".docx": _read_docx_full,
    ".csv": _read_csv_full,
    ".json": _read_json_full,
}


def _classify_type(ext: str) -> str:
    """与 library_store._classify 同语义；此处内置一份避免跨包私有依赖。"""
    ext = ext.lower()
    if ext in (".md", ".txt"):
        return "text"
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext == ".csv":
        return "csv"
    if ext == ".json":
        return "json"
    return "other"


# ────────────────────────────────────────────────────────────────
# 公开 API
# ────────────────────────────────────────────────────────────────

def read_file_full(path: Path | str) -> dict:
    """**完整**读取单文件文本（不限字符 / 不限页数）。

    返回:
        {
          "path": str,        # 绝对路径（字符串）
          "name": str,        # basename
          "type": str,        # "text" / "pdf" / "docx" / "csv" / "json" / "other"
          "text": str,        # 完整正文；失败/空为 ""
          "ok": bool,         # 是否成功提取到非空文本
          "chars": int,       # text 字符数（len(text)）
        }

    失败/缺库/不存在：ok=False, text=""，**绝不抛异常**。
    """
    p = Path(path) if not isinstance(path, Path) else path
    result = {
        "path": str(p),
        "name": p.name,
        "type": _classify_type(p.suffix),
        "text": "",
        "ok": False,
        "chars": 0,
    }
    if not p.exists() or not p.is_file():
        log.warning("[readers] 文件不存在或不是文件: %s", p)
        return result
    try:
        ext = p.suffix.lower()
        fn = _FULL_EXTRACTORS.get(ext, _read_text_file_full)
        text = fn(p) or ""
        result["text"] = text
        result["chars"] = len(text)
        result["ok"] = bool(text.strip())
        return result
    except Exception as e:
        log.warning("[readers] read_file_full 顶层异常 %s: %s", p, e)
        return result


def read_corpus_full(learner_id: str) -> List[dict]:
    """读取某个 learner 资料库的**全部**文件全文。

    内部走 library_store.list_user_files（自动双读 usr_knowledge + 旧路径）。
    返回 list[{path, name, type, text, ok, chars}]；
    无文件返回 []；单文件失败不影响整体。
    """
    if not learner_id:
        learner_id = "anonymous"
    try:
        paths = library_store.list_user_files(learner_id)
    except Exception as e:
        log.warning("[readers] list_user_files 失败 learner=%s: %s", learner_id, e)
        return []
    out: List[dict] = []
    for p in paths:
        out.append(read_file_full(p))
    return out


__all__ = ["read_file_full", "read_corpus_full"]